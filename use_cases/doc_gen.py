"""Document generation for EmberOS-Windows.

Supports writing content to TXT, Markdown, PDF, and DOCX formats.
PDF uses fpdf2 and DOCX uses python-docx — both are optional with
graceful fallbacks to plain text when not installed.
"""

import logging
from pathlib import Path

logger = logging.getLogger("emberos.use_cases.doc_gen")


def write_document(path: str, content: str, fmt: str = None) -> str:
    """Write content to a file.

    Format is determined by the file extension or the fmt override.
    Supported: txt, md, pdf, docx.
    Falls back to plain text when a format-specific library is absent.
    """
    p = Path(path)
    ext = (fmt or p.suffix or ".txt").lstrip(".").lower()
    if ext == "markdown":
        ext = "md"

    p.parent.mkdir(parents=True, exist_ok=True)

    if ext in ("txt", "md"):
        return _write_text(p, content)
    elif ext == "pdf":
        return _write_pdf(p, content)
    elif ext == "docx":
        return _write_docx(p, content)
    else:
        return _write_text(p.with_suffix(".txt"), content)


def _write_text(path: Path, content: str) -> str:
    try:
        path.write_text(content, encoding="utf-8")
        size = path.stat().st_size
        size_str = f"{size // 1024} KB" if size >= 1024 else f"{size} B"
        return f"Document saved: {path}  ({size_str})"
    except Exception as e:
        return f"Failed to write document: {e}"


def _write_pdf(path: Path, content: str) -> str:
    """Write a PDF using fpdf2. Falls back to .txt if fpdf2 is absent."""
    try:
        from fpdf import FPDF  # noqa: F401

        pdf = FPDF()
        pdf.set_margins(20, 20, 20)
        pdf.add_page()

        for line in content.split("\n"):
            stripped = line.strip()
            if stripped.startswith("### "):
                pdf.set_font("Helvetica", "B", size=12)
                pdf.multi_cell(0, 8, stripped[4:] or " ")
                pdf.set_font("Helvetica", size=11)
            elif stripped.startswith("## "):
                pdf.set_font("Helvetica", "B", size=13)
                pdf.multi_cell(0, 9, stripped[3:] or " ")
                pdf.set_font("Helvetica", size=11)
            elif stripped.startswith("# "):
                pdf.set_font("Helvetica", "B", size=15)
                pdf.multi_cell(0, 11, stripped[2:] or " ")
                pdf.set_font("Helvetica", size=11)
            else:
                pdf.set_font("Helvetica", size=11)
                pdf.multi_cell(0, 7, line if line.strip() else " ")

        pdf.output(str(path))
        size = path.stat().st_size
        size_str = f"{size // 1024} KB" if size >= 1024 else f"{size} B"
        return f"PDF saved: {path}  ({size_str})"

    except ImportError:
        txt_path = path.with_suffix(".txt")
        try:
            txt_path.write_text(content, encoding="utf-8")
            return (
                f"fpdf2 is not installed — saved as plain text: {txt_path}\n"
                "To enable PDF output run:  "
                "env\\python-embed\\python.exe -m pip install fpdf2"
            )
        except Exception as e:
            return f"Failed to write document: {e}"
    except Exception as e:
        return f"PDF generation failed: {e}"


def _write_docx(path: Path, content: str) -> str:
    """Write a DOCX using python-docx. Falls back to .txt if absent."""
    try:
        from docx import Document  # noqa: F401

        doc = Document()
        for line in content.split("\n"):
            stripped = line.strip()
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)

        doc.save(str(path))
        size = path.stat().st_size
        size_str = f"{size // 1024} KB" if size >= 1024 else f"{size} B"
        return f"DOCX saved: {path}  ({size_str})"

    except ImportError:
        txt_path = path.with_suffix(".txt")
        try:
            txt_path.write_text(content, encoding="utf-8")
            return (
                f"python-docx is not installed — saved as plain text: {txt_path}\n"
                "To enable DOCX output run:  "
                "env\\python-embed\\python.exe -m pip install python-docx"
            )
        except Exception as e:
            return f"Failed to write document: {e}"
    except Exception as e:
        return f"DOCX generation failed: {e}"
